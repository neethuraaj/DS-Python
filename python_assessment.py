import os
import shutil
import pandas as pd
import pdfplumber
from pdf2docx import Converter
from docx import Document
from openpyxl.styles import Alignment, PatternFill, Border, Side, Font
import tkinter as tk
from tkinter import filedialog, messagebox
import subprocess

pdf_path = ""
excel_path = ""

# ------------------ TASK 1 ------------------
def task1(excel_file, pdf_file):
    # 1. Read Excel
    table1 = pd.read_excel(excel_file, sheet_name="Employee-Department Data")

    # 2. Read PDF
    with pdfplumber.open(pdf_file) as pdf:
        page = pdf.pages[0]
        table2_raw = page.extract_table()
    table2 = pd.DataFrame(table2_raw[1:], columns=table2_raw[0])

    # 3. Join tables
    final_table = pd.merge(table1, table2, on="Employee ID")
    final_table = final_table[['Employee ID','Name','Email','Department','Joining Date',
                               'Project Name','Start Date','End Date','Allocations (%)']]

    # 4. Date conversion + new columns
    final_table['Joining Date'] = pd.to_datetime(final_table['Joining Date'], dayfirst=True)
    final_table['Start Date'] = pd.to_datetime(final_table['Start Date'], dayfirst=True)
    final_table['End Date'] = pd.to_datetime(final_table['End Date'], dayfirst=True)

    final_table['Project Duration (Days)'] = (final_table['End Date'] - final_table['Start Date']).dt.days
    final_table['Bench Duration'] = (final_table['Start Date'] - final_table['Joining Date']).dt.days

    # 5. Correct/Wrong tables
    wrong_table = final_table[final_table['Project Duration (Days)'] < 0]
    correct_table = final_table[final_table['Project Duration (Days)'] >= 0]

    # 6. Save with formatting
    with pd.ExcelWriter("Output_File.xlsx", engine="openpyxl") as writer:
        final_table.to_excel(writer, sheet_name="Final Data", index=False)
        correct_table.to_excel(writer, sheet_name="Correct Data", index=False)
        wrong_table.to_excel(writer, sheet_name="Wrong Data", index=False)

        wb = writer.book
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            # Header formatting
            for cell in ws[1]:
                cell.alignment = Alignment(horizontal="center")
                cell.fill = PatternFill("solid", fgColor="90EE90")
                cell.font = Font(color="000000", bold=True)
            # Cell formatting
            for row in ws.iter_rows():
                for cell in row:
                    cell.alignment = Alignment(horizontal="center")
                    cell.border = Border(left=Side(style="thin"), right=Side(style="thin"),
                                         top=Side(style="thin"), bottom=Side(style="thin"))
            ws.sheet_view.showGridLines = False
        wb.save("Output_File.xlsx")

    return final_table, table1


# ------------------ TASK 2 ------------------
def task2(final_table, table1, pdf_file):
    # 1. Convert PDF → Word
    cv = Converter(pdf_file)
    cv.convert("Project_Allocation_Report.docx", start=0, end=None)
    cv.close()

    # 2. Department counts
    dept_counts = final_table['Department'].value_counts()
    highest = dept_counts.index[0]
    second_highest = dept_counts.index[1]

    # 3. Replace placeholders in Word
    doc = Document("Project_Allocation_Report.docx")
    for p in doc.paragraphs:
        if "[highest]" in p.text:
            p.text = p.text.replace("[highest]", highest)
        if "[second highest]" in p.text:
            p.text = p.text.replace("[second highest]", second_highest)
        if "<employee_table_123>" in p.text:
            p.text = p.text.replace("<employee_table_123>", "")
            table = doc.add_table(rows=1, cols=len(table1.columns))
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, col in enumerate(table1.columns):
                hdr[i].text = col
            for idx, row in table1.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)

    doc.save("Project_Allocation_Report_Output.docx")

    # 4. Move files into Output folder
    os.makedirs("Output", exist_ok=True)
    shutil.move("Output_File.xlsx", "Output/Output_File.xlsx")
    shutil.move("Project_Allocation_Report_Output.docx", "Output/Project_Allocation_Report_Output.docx")

    return os.path.abspath("Output")


# ------------------ TASK 3 (UI) ------------------
def run_all():
    global pdf_path, excel_path
    if not pdf_path or not excel_path:
        messagebox.showerror("Error", "Please select both PDF and Excel files")
        return

    final_table, table1 = task1(excel_path, pdf_path)
    output_folder = task2(final_table, table1, pdf_path)

    messagebox.showinfo("Success", f"All tasks completed successfully! Files saved in {output_folder}")

    # Open the output folder automatically
    if os.name == 'nt':  # Windows
        os.startfile(output_folder)
    elif os.name == 'posix':  # macOS/Linux
        subprocess.call(["open" if sys.platform == "darwin" else "xdg-open", output_folder])


def browse_pdf():
    global pdf_path
    pdf_path = filedialog.askopenfilename(filetypes=[("PDF files","*.pdf")])
    messagebox.showinfo("Selected", f"PDF: {pdf_path}")


def browse_excel():
    global excel_path
    excel_path = filedialog.askopenfilename(filetypes=[("Excel files","*.xlsx")])
    messagebox.showinfo("Selected", f"Excel: {excel_path}")


def start_ui():
    root = tk.Tk()
    root.title("Python Assessment")
    root.geometry("350x250")

    tk.Label(root, text="Python Assessment Tool", font=("Arial", 14, "bold")).pack(pady=10)
    tk.Button(root, text="Browse PDF", command=browse_pdf).pack(pady=5)
    tk.Button(root, text="Browse Excel", command=browse_excel).pack(pady=5)
    tk.Button(root, text="Run Assessment", command=run_all).pack(pady=20)

    root.mainloop()


# ------------------ MAIN ------------------
if __name__ == "__main__":
    start_ui()
